import streamlit as st
st.set_page_config(page_title='Audio Transcription Parsing')
st.header('Audio Transcription Parsing')
video = st.file_uploader("Upload the video file", type=["mp4"])
transcription = st.file_uploader("Upload the transcription file", type=["vtt"])
presentation = st.file_uploader("Upload the presentation file(as a PDF)", type=["pdf"])

from pdf2image import convert_from_path
import cv2
import numpy as np
import os
import webvtt
import json
import requests
from PIL import Image
import base64
from io import BytesIO
from docx import Document
import tempfile

ENDPOINT = "https://gpt-4omniwithimages.openai.azure.com/openai/deployments/GPT4Omni/chat/completions?api-version=2024-02-15-preview"
API_KEY = "6e98566acaf24997baa39039b6e6d183"


def extract_slides_from_pdf(pdf_file):
    slides = convert_from_path(pdf_file)
    slide_images = []

    for slide in slides:
        slide_array = np.array(slide)
        slide_images.append(slide_array)
    
    return slide_images

def parse_vtt(vtt_file):
    segments = []
    for caption in webvtt.read(vtt_file):
        segments.append({
            'start': caption.start,
            'end': caption.end,
            'text': caption.text
        })
    return segments

def detect_slide_changes(video_path, threshold=30, min_contour_area=5000):
    cap = cv2.VideoCapture(video_path)
    frame_rate = cap.get(cv2.CAP_PROP_FPS)
    slide_times = []

    ret, prev_frame = cap.read()
    prev_frame_gray = cv2.cvtColor(prev_frame, cv2.COLOR_BGR2GRAY)
    frame_count = 0

    while ret:
        ret, frame = cap.read()
        if not ret:
            break

        frame_gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        frame_diff = cv2.absdiff(prev_frame_gray, frame_gray)
        _, thresh = cv2.threshold(frame_diff, threshold, 255, cv2.THRESH_BINARY)
        contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        if any(cv2.contourArea(contour) > min_contour_area for contour in contours):
            slide_time = frame_count / frame_rate
            slide_times.append(slide_time)
            prev_frame_gray = frame_gray

        frame_count += 1

    cap.release()
    return slide_times

def format_conversion(time):
    hour=int(time)//3600
    minute=int(time)//60
    second=time%60
    second=round(second,3)
    corrected_fomrat=f"{hour:02d}:{minute:02d}:{second:06.3f}"
    return corrected_fomrat

def extract_transcript_for_slide(transcript, start_time, end_time):
    """Extract transcript segments that fall within the specified time range."""
    start_time_seconds = format_conversion(start_time)
    end_time_seconds = format_conversion(end_time)

    slide_transcript = [segment['text'] for segment in transcript if segment['start'] >= start_time_seconds and segment['end'] <= end_time_seconds]
    return " ".join(slide_transcript)

def find_midpoints(slide_times):
    midpoints = [(slide_times[i] + slide_times[i+1]) / 2 for i in range(len(slide_times) - 1)]
    return midpoints


def extract_frames_at_midpoints(video_file, midpoints, crop_fraction=1.0):
    """
    Extracts frames from the video at the given midpoints and crops the central region of each frame.

    Args:
        video_file (str): Path to the video file.
        midpoints (list): List of midpoint timestamps (in seconds) where frames should be extracted.
        crop_fraction (float): Fraction of the frame to retain as the central region (0 < crop_fraction <= 1).

    Returns:
        list: List of cropped frames in RGB format.
    """
    cap = cv2.VideoCapture(video_file)
    frames = []

    for midpoint in midpoints:
        # Set video position to the midpoint timestamp
        cap.set(cv2.CAP_PROP_POS_MSEC, midpoint * 1000)
        ret, frame = cap.read()
        if ret:
            # Crop the central region of the frame
            height, width = frame.shape[:2]  # Only get height and width, ignore channels
            crop_height = int(height * crop_fraction)
            crop_width = int(width * crop_fraction)
            start_x = (width - crop_width) // 2
            start_y = (height - crop_height) // 2
            cropped_frame = frame[start_y:start_y + crop_height, start_x:start_x + crop_width]

            # Convert from BGR to RGB
            rgb_frame = cv2.cvtColor(cropped_frame, cv2.COLOR_BGR2RGB)
            rgb_frame = cv2.resize(rgb_frame, (2667, 1500))
            frames.append(rgb_frame)

    cap.release()
    return [frames,midpoints]

def gpt_response(encoded_image1,encoded_image2):
    headers = {
    "Content-Type": "application/json",
    "api-key": API_KEY,
    }

    # Payload for the request
    payload = {
    "messages": [
        {
        "role": "system",
        "content": [
            {
            "type": "text",
            "text": "Compare the 2 images. If the contents are similar, return 'True'. If not, return 'False'. Always remember that the output you give must only be either True or False."
            }
        ]
        },
        {
        "role": "user",
        "content": [
            {
            "type": "image",
            "image": encoded_image1
            }
        ]
        },
        {
        "role": "user",
        "content": [
            {
            "type": "image",
            "image": encoded_image2
            }
        ]
        }
    ],
    "temperature": 0.2,
    "max_tokens": 20
    }
    response = requests.post(ENDPOINT, headers=headers, json=payload)
    return response.json()['choices'][0]['message']['content']

def create_word_document(data):
    doc = Document()
    for key, content in data.items():
        doc.add_heading(f'Slide {key + 1}', level=1)
        doc.add_paragraph(content)
    return doc

def save_uploaded_file(uploaded_file):
    """
    Save the uploaded file to a temporary directory and return the file path and file name.
    
    Args:
        uploaded_file: The uploaded file.

    Returns:
        tuple: The path to the saved file and the original file name.
    """
    file_extension = os.path.splitext(uploaded_file.name)[1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
        tmp_file.write(uploaded_file.read())
        return tmp_file.name

def format_text(transcript_mapping):
    for key, value in transcript_mapping.items():
        if len(value)>0:
            headers = {
                "Content-Type": "application/json",
                "api-key": API_KEY,
            }

            # Payload for the request
            payload = {
                "messages": [
                    {
                        "role": "system",
                        "content": "You are given a text which is a section of the audio transcriptions corresponding to a video. Your task is to format the information from the text in the form of bullet points. Keep in mind that all the keywords, important terminology, and key information should be included in your response. Also, use only the information from the text provided to you for generating bullet points."
                    },
                    {
                        "role": "user",
                        "content": value  # Pass the text from the transcript mapping
                    },
                ],
                "temperature": 0.2,
                "max_tokens": 800
            }

            response = requests.post(ENDPOINT, headers=headers, json=payload)
            transcript_mapping[key] = response.json()['choices'][0]['message']['content']  # Save the formatted text back into the mapping

    return transcript_mapping

if (video and transcription and presentation):
    try:
        saved_presentation_path = save_uploaded_file(presentation)
        slide_images = extract_slides_from_pdf(saved_presentation_path)
        
        saved_video_path = save_uploaded_file(video)
        saved_transcription_path = save_uploaded_file(transcription)
        
        transcripts = parse_vtt(saved_transcription_path)
        slide_times = detect_slide_changes(saved_video_path)
            


    # video_file=video.name
    # transcription_file=transcription.name
    # # presentation_file=presentation.name
    # # slide_images = extract_slides_from_pdf(presentation_file)
    # transcripts = parse_vtt(transcription_file)
    # slide_times=detect_slide_changes(video_file)
        groups = []
        current_group = [slide_times[0]]
        for i in range(1, len(slide_times)):
            if slide_times[i] - current_group[-1] <= 5:
                current_group.append(slide_times[i])
            else:
                groups.append(current_group)
                current_group = [slide_times[i]]

        if current_group:
            groups.append(current_group)

        middle_values = [0.0]
        for group in groups:
            middle = sum(group) / len(group)
            middle_values.append(middle)
        cap = cv2.VideoCapture(saved_video_path)
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        end_time = total_frames / fps
        middle_values.append(end_time)
        midpoints=find_midpoints(middle_values)
        frames,midpoints=extract_frames_at_midpoints(video_file=saved_video_path,midpoints=midpoints,crop_fraction=1.0)
        list=[]
        for i in range(len(middle_values)-1):
            list.append([middle_values[i],middle_values[i+1]])
        transcript_align=[]
        for frame, time_stamps in zip(frames,list):
            transcript_align.append([frame,time_stamps])
        transcript_mapping = {i: "" for i in range(len(slide_images))}

        for i in range(len(transcript_align)):
            for j in range(len(slide_images)):
                frame = transcript_align[i][0]
                frame_image = Image.fromarray(frame)
                slide_image = Image.fromarray(slide_images[j])

                # Convert frame image to base64
                buffer1 = BytesIO()
                frame_image.save(buffer1, format="PNG")
                buffer1.seek(0)
                frame_data = buffer1.read()
                encoded_image1 = base64.b64encode(frame_data).decode('utf-8')

                # Convert slide image to base64
                buffer2 = BytesIO()
                slide_image.save(buffer2, format="PNG")
                buffer2.seek(0)
                slide_data = buffer2.read()
                encoded_image2 = base64.b64encode(slide_data).decode('utf-8')

                # Get the response from GPT
                response = gpt_response(encoded_image1, encoded_image2)
                print(response)

                if response.strip()[:4] in('true','True','True.','true.','TRUE','TRUE.'):
                    ans = j
                    answer = extract_transcript_for_slide(transcripts, transcript_align[i][1][0], transcript_align[i][1][1])
                    if transcript_mapping[ans]:
                        transcript_mapping[ans] += ' ' + answer
                    else:
                        transcript_mapping[ans] = answer
                    break

        formatted_mapping=format_text(transcript_mapping)
        for key, content in formatted_mapping.items():
            st.write(f'Slide {key + 1}' )
            st.write('\n',content,'\n')

    except Exception as e:
        st.error(f"An error occurred: {e}")

    doc = create_word_document(formatted_mapping)
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    st.download_button(
    label="Download Word Document",
    data=doc_io,
    file_name='Slides_Document.docx',
    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )